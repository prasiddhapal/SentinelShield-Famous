"""
SentinelShield — Advanced Intrusion Detection & Web Protection System
WAF + IDS with SQLite persistent database backend.
"""

from flask import Flask, render_template, request, jsonify, g, Response
import sqlite3
from datetime import datetime
import threading
import time
import requests
import re
from fpdf import FPDF
from docx import Document
import csv
import io
import os
import sqlite3
import threading
from datetime import datetime, timedelta
from collections import defaultdict

app = Flask(__name__)

# ─────────────────────────────────────────────
#  Paths
# ─────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH  = os.path.join(BASE_DIR, "sentinel.db")
LOG_DIR  = os.path.join(BASE_DIR, "logs")
os.makedirs(LOG_DIR, exist_ok=True)

# ─────────────────────────────────────────────
#  Config
# ─────────────────────────────────────────────
RATE_LIMIT_MAX    = 10   # max requests per IP per window
RATE_LIMIT_WINDOW = 30   # seconds

# ─────────────────────────────────────────────
#  Database Setup
# ─────────────────────────────────────────────
db_lock = threading.Lock()

def get_db():
    """Open a thread-local DB connection."""
    if not hasattr(g, "_db"):
        g._db = sqlite3.connect(DB_PATH)
        g._db.row_factory = sqlite3.Row
    return g._db

@app.teardown_appcontext
def close_db(exc):
    db = getattr(g, "_db", None)
    if db:
        db.close()

def init_db():
    """Create tables if they don't exist, safely migrate columns."""
    with sqlite3.connect(DB_PATH) as conn:
        conn.executescript("""
        CREATE TABLE IF NOT EXISTS attack_logs (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            timestamp   TEXT    NOT NULL,
            ip          TEXT    NOT NULL,
            method      TEXT    NOT NULL,
            path        TEXT    NOT NULL,
            query       TEXT    DEFAULT '',
            attack_type TEXT    NOT NULL,
            severity    TEXT    NOT NULL,
            action      TEXT    NOT NULL,
            payload     TEXT    DEFAULT '',
            risk_score  INTEGER DEFAULT 0,
            risk_level  TEXT    DEFAULT 'low'
        );

        CREATE INDEX IF NOT EXISTS idx_logs_timestamp   ON attack_logs(timestamp);
        CREATE INDEX IF NOT EXISTS idx_logs_ip          ON attack_logs(ip);
        CREATE INDEX IF NOT EXISTS idx_logs_action      ON attack_logs(action);
        CREATE INDEX IF NOT EXISTS idx_logs_attack_type ON attack_logs(attack_type);

        CREATE TABLE IF NOT EXISTS ip_reputation (
            ip              TEXT PRIMARY KEY,
            total_requests  INTEGER DEFAULT 0,
            blocked_count   INTEGER DEFAULT 0,
            flagged_count   INTEGER DEFAULT 0,
            allowed_count   INTEGER DEFAULT 0,
            first_seen      TEXT,
            last_seen       TEXT
        );

        CREATE TABLE IF NOT EXISTS banned_ips (
            ip              TEXT PRIMARY KEY,
            reason          TEXT,
            timestamp       TEXT,
            total_risk_at_ban INTEGER
        );

        CREATE TABLE IF NOT EXISTS ip_geo_cache (
            ip              TEXT PRIMARY KEY,
            country_code    TEXT,
            country_name    TEXT,
            latitude        REAL,
            longitude       REAL
        );
        """)

        # Safely add new columns to existing DBs (ignored if already present)
        for col_sql in [
            "ALTER TABLE attack_logs ADD COLUMN risk_score INTEGER DEFAULT 0",
            "ALTER TABLE attack_logs ADD COLUMN risk_level TEXT DEFAULT 'low'",
            "ALTER TABLE attack_logs ADD COLUMN country_code TEXT DEFAULT 'US'",
            "ALTER TABLE attack_logs ADD COLUMN country_name TEXT DEFAULT 'United States'",
        ]:
            try:
                conn.execute(col_sql)
            except sqlite3.OperationalError:
                pass  # Column already exists

        # Add risk_level index after ensuring column exists
        try:
            conn.execute("CREATE INDEX IF NOT EXISTS idx_logs_risk_level ON attack_logs(risk_level)")
        except sqlite3.OperationalError:
            pass

        conn.commit()
    print(f"[DB] Initialized: {DB_PATH}")



# ─────────────────────────────────────────────
#  Risk Scoring Engine
# ─────────────────────────────────────────────

# Base risk scores per attack type (0–100 scale)
ATTACK_BASE_SCORES = {
    'Command Injection':          95,
    'Insecure Deserialization':   92,
    'Local File Inclusion (LFI)': 90,
    'SQL Injection':              85,
    'XML External Entity (XXE)':  80,
    'SSRF Attack':                75,
    'Cross-Site Scripting (XSS)': 70,
    'Directory Traversal':        65,
    'Brute-force / Rate Limit':   55,
    'Normal':                      5,
}

SEVERITY_BONUS = {
    'critical': 5,
    'high':     3,
    'medium':   1,
    'info':     0,
}

def risk_level_label(score: int) -> str:
    if score >= 76: return 'critical'
    if score >= 51: return 'high'
    if score >= 26: return 'medium'
    return 'low'

def calculate_risk_score(attack_type: str, severity: str, action: str, ip: str) -> int:
    """Compute 0–100 risk score for a single request."""
    base = ATTACK_BASE_SCORES.get(attack_type, 10)
    bonus = SEVERITY_BONUS.get(severity, 0)

    # Scale by action
    if action == 'ALLOWED':
        score = max(int(base * 0.15) + bonus, 2)  # Low but not zero
    elif action == 'FLAGGED':
        score = int(base * 0.75) + bonus
    else:  # BLOCKED
        score = base + bonus

    # Repeat-offender bonus: check in-memory rate tracker
    with rate_lock:
        hits = len(rate_tracker.get(ip, []))
    if hits > RATE_LIMIT_MAX:
        score = min(score + 10, 100)
    elif hits > 5:
        score = min(score + 5, 100)

    return min(score, 100)

def get_geo_data(ip: str):
    """
    GeoIP lookup with SQLite caching. Non-blocking fallback on timeout.
    """
    if ip in ("127.0.0.1", "::1", "0.0.0.0", "localhost"):
        import random
        SIM_COUNTRIES = [
            ("RU", "Russia"), ("CN", "China"), ("US", "United States"),
            ("BR", "Brazil"), ("NL", "Netherlands"), ("DE", "Germany"),
            ("IN", "India"), ("KP", "North Korea"), ("UA", "Ukraine")
        ]
        return random.choice(SIM_COUNTRIES)

    # 1. Check sqlite cache first (fast, no I/O wait)
    try:
        with sqlite3.connect(DB_PATH, timeout=2) as conn:
            conn.row_factory = sqlite3.Row
            cached = conn.execute("SELECT country_code, country_name FROM ip_geo_cache WHERE ip = ?", (ip,)).fetchone()
            if cached:
                return cached["country_code"], cached["country_name"]
    except Exception:
        pass

    # 2. Not cached — fetch from API with strict timeout
    try:
        r = requests.get(f"http://ip-api.com/json/{ip}?fields=status,country,countryCode,lat,lon", timeout=2)
        if r.status_code == 200:
            data = r.json()
            if data.get("status") == "success":
                cc = data.get("countryCode", "??")
                cn = data.get("country", "Unknown")
                lat = data.get("lat", 0.0)
                lon = data.get("lon", 0.0)
                with sqlite3.connect(DB_PATH, timeout=3) as conn:
                    conn.execute("""
                        INSERT OR IGNORE INTO ip_geo_cache (ip, country_code, country_name, latitude, longitude)
                        VALUES (?, ?, ?, ?, ?)
                    """, (ip, cc, cn, lat, lon))
                    conn.commit()
                return cc, cn
    except Exception as e:
        print(f"[GeoIP] Fetch Error for {ip}: {e}")

    return ("??", "Unknown")


def write_log(entry: dict):
    """Insert log entry into SQLite and update ip_reputation."""
    with db_lock:
        try:
            with sqlite3.connect(DB_PATH) as conn:
                conn.row_factory = sqlite3.Row
                conn.execute("""
                    INSERT INTO attack_logs
                        (timestamp, ip, method, path, query, attack_type, severity, action, payload, risk_score, risk_level, country_code, country_name)
                    VALUES
                        (:timestamp, :ip, :method, :path, :query, :attack_type, :severity, :action, :payload, :risk_score, :risk_level, :country_code, :country_name)
                """, entry)

                # Upsert ip_reputation
                conn.execute("""
                    INSERT INTO ip_reputation (ip, total_requests, blocked_count, flagged_count, allowed_count, first_seen, last_seen)
                    VALUES (:ip, 1,
                        CASE WHEN :action='BLOCKED' THEN 1 ELSE 0 END,
                        CASE WHEN :action='FLAGGED' THEN 1 ELSE 0 END,
                        CASE WHEN :action='ALLOWED' THEN 1 ELSE 0 END,
                        :timestamp, :timestamp)
                    ON CONFLICT(ip) DO UPDATE SET
                        total_requests = total_requests + 1,
                        blocked_count  = blocked_count  + CASE WHEN :action='BLOCKED' THEN 1 ELSE 0 END,
                        flagged_count  = flagged_count  + CASE WHEN :action='FLAGGED' THEN 1 ELSE 0 END,
                        allowed_count  = allowed_count  + CASE WHEN :action='ALLOWED' THEN 1 ELSE 0 END,
                        last_seen      = :timestamp
                """, entry)
                
                # Check for Auto-Ban Threshold (200 cumulative risk)
                res = conn.execute("SELECT SUM(risk_score) as total_risk FROM attack_logs WHERE ip = ?", (entry["ip"],)).fetchone()
                total_risk = (res["total_risk"] or 0) if res else 0
                
                if total_risk >= 200:
                    conn.execute("""
                        INSERT OR IGNORE INTO banned_ips (ip, reason, timestamp, total_risk_at_ban)
                        VALUES (?, ?, ?, ?)
                    """, (entry["ip"], "Exceeded Security Risk Threshold (200+)", entry["timestamp"], total_risk))
                
                conn.commit()
        except Exception as e:
            print(f"[DB] Write error: {e}")

# ─────────────────────────────────────────────
#  Rate Limiter (in-memory, fast)
# ─────────────────────────────────────────────
rate_tracker = defaultdict(list)
rate_lock    = threading.Lock()

def check_rate_limit(ip: str) -> bool:
    now = time.time()
    with rate_lock:
        rate_tracker[ip] = [t for t in rate_tracker[ip] if now - t < RATE_LIMIT_WINDOW]
        rate_tracker[ip].append(now)
        return len(rate_tracker[ip]) > RATE_LIMIT_MAX

def rate_hit_count(ip: str) -> int:
    with rate_lock:
        return len(rate_tracker.get(ip, []))

# ─────────────────────────────────────────────
#  Attack Signature Rules
# ─────────────────────────────────────────────
ATTACK_RULES = [
    {
        "name": "SQL Injection",
        "severity": "critical",
        "patterns": [
            r"(?i)(\bOR\b|\bAND\b)\s+['\"]?\d+['\"]?\s*=\s*['\"]?\d+",
            r"(?i)'?\s*(OR|AND)\s+'?[a-z0-9]+'?\s*=\s*'?[a-z0-9]+",
            r"(?i)\bDROP\s+TABLE\b",
            r"(?i)\bINSERT\s+INTO\b.*\bVALUES\b",
            r"(?i)\bUNION\s+(ALL\s+)?SELECT\b",
            r"(?i)\bSELECT\b.*\bFROM\b.*\bWHERE\b",
            r"(?i)--\s*$",
            r"(?i)['\"];\s*(DROP|DELETE|INSERT|UPDATE|SELECT)\b",
            r"(?i)\bEXEC(\s+|\()",
            r"(?i)xp_cmdshell",
            r"(?i)\bSLEEP\s*\(\s*\d+",
            r"(?i)\bBENCHMARK\s*\(",
        ],
    },
    {
        "name": "Cross-Site Scripting (XSS)",
        "severity": "high",
        "patterns": [
            r"(?i)<\s*script[^>]*>",
            r"(?i)</\s*script\s*>",
            r"(?i)\bon\w+\s*=",
            r"(?i)javascript\s*:",
            r"(?i)<\s*iframe[^>]*>",
            r"(?i)<\s*img[^>]+onerror",
            r"(?i)document\.cookie",
            r"(?i)alert\s*\(",
            r"(?i)eval\s*\(",
            r"(?i)<\s*svg[^>]*onload",
        ],
    },
    {
        "name": "Directory Traversal",
        "severity": "high",
        "patterns": [
            r"\.\./",
            r"\.\.\\",
            r"(?i)%2e%2e%2f",
            r"(?i)%2e%2e/",
            r"(?i)\.\.%2f",
        ],
    },
    {
        "name": "Local File Inclusion (LFI)",
        "severity": "critical",
        "patterns": [
            r"(?i)/etc/(passwd|shadow|hosts|group|fstab|sudoers)",
            r"(?i)/proc/(self|version|cmdline|environ)",
            r"(?i)/var/log/",
            r"(?i)C:\\Windows\\System32",
            r"(?i)win\.ini",
            r"(?i)boot\.ini",
        ],
    },
    {
        "name": "Command Injection",
        "severity": "critical",
        "patterns": [
            r"(?i);\s*(ls|cat|rm|wget|curl|chmod|chown|kill|ps|whoami|id|uname|bash|sh|python|perl|ruby|php)\b",
            r"(?i)\|\s*(ls|cat|rm|wget|curl|chmod|id|whoami|uname|bash|sh)\b",
            r"(?i)&&\s*(ls|cat|rm|wget|curl|chmod|id|whoami|uname|bash|sh)\b",
            r"(?i)\$\(.*\)",
            r"(?i)`.*`",
            r"(?i)\bwhoami\b",
            r"(?i)\b(cmd\.exe|/bin/bash|/bin/sh)\b",
            r"(?i);cat\s+/",
        ],
    },
    {
        "name": "SSRF Attack",
        "severity": "high",
        "patterns": [
            r"(?i)\b169\.254\.169\.254\b",
            r"(?i)\bmetadata\.google\.internal\b",
            r"(?i)\burl=(http|https)://(localhost|127\.0\.0\.1|0\.0\.0\.0)",
            r"(?i)(fetch|proxy|curl|wget).+?(localhost|127\.0\.0\.1)",
        ],
    },
    {
        "name": "XML External Entity (XXE)",
        "severity": "critical",
        "patterns": [
            r"(?i)<!ENTITY\s+[%a-zA-Z0-9_]+\s+SYSTEM\s+",
            r"(?i)<!DOCTYPE\s+\w+\s+\[\s*<!ENTITY",
            r"(?i)(file|http|ftp|gopher)://",
            r"(?i)SYSTEM\s+['\"][^'\"]+['\"]",
        ],
    },
    {
        "name": "Insecure Deserialization",
        "severity": "critical",
        "patterns": [
            r"(?i)O:\d+:\s*['\"][^'\"]*['\"]:\d+:",
            r"(?i)a:\d+:\{",
            r"(?i)c__main__\s*\\n",
            r"(?i)x0c__main__",
            r"(?i)cos\s*\\nsystem",
            r"(?i)rO0AB",  # Base64 start of Java serialized object
        ],
    },
]

# ─────────────────────────────────────────────
#  Core WAF Functions
# ─────────────────────────────────────────────

def get_client_ip():
    return request.headers.get("X-Forwarded-For", request.remote_addr)

def inspect_payload(payload: str):
    if not payload:
        return None, None, None
    for rule in ATTACK_RULES:
        for pattern in rule["patterns"]:
            if re.search(pattern, payload):
                return rule["name"], rule["severity"], payload[:120].strip()
    return None, None, None

def build_payload_string() -> str:
    parts = [request.path, request.query_string.decode("utf-8", errors="ignore")]
    for h in ("User-Agent", "Referer", "Cookie"):
        v = request.headers.get(h, "")
        if v:
            parts.append(v)
    try:
        body = request.get_data(as_text=True)
        if body:
            parts.append(body)
    except Exception:
        pass
    return " ".join(parts)

# ─────────────────────────────────────────────
#  Before-request WAF Middleware
# ─────────────────────────────────────────────

@app.before_request
def waf_middleware():
    if request.path.startswith("/static") or request.path.startswith("/api"):
        return

    ip = get_client_ip()

    # 0. Check for Permanent Ban
    try:
        with sqlite3.connect(DB_PATH, timeout=3) as conn:
            conn.row_factory = sqlite3.Row
            ban = conn.execute("SELECT * FROM banned_ips WHERE ip = ?", (ip,)).fetchone()
            if ban:
                return render_template("banned.html", ip=ip, timestamp=ban["timestamp"]), 403
    except Exception:
        pass

    payload = build_payload_string()

    # 1. Signature matching
    attack_type, severity, payload_excerpt = inspect_payload(payload)
    if attack_type:
        action = "BLOCKED"
    elif check_rate_limit(ip):
        hits        = rate_hit_count(ip)
        action      = "FLAGGED" if hits == RATE_LIMIT_MAX + 1 else "BLOCKED"
        attack_type = "Brute-force / Rate Limit"
        severity    = "medium"
        payload_excerpt = request.path[:80]
    else:
        action          = "ALLOWED"
        attack_type     = "Normal"
        severity        = "info"
        payload_excerpt = request.path[:80]

    # 2. Compute risk score immediately
    score = calculate_risk_score(attack_type, severity, action, ip)
    level = risk_level_label(score)

    # 3. Geo — instant cache check only, schedule API call in background
    cc, cn = "??", "Unknown"
    if ip in ("127.0.0.1", "::1", "0.0.0.0", "localhost"):
        import random
        cc, cn = random.choice([
            ("RU", "Russia"), ("CN", "China"), ("US", "United States"),
            ("BR", "Brazil"), ("DE", "Germany"), ("IN", "India"),
            ("KP", "North Korea"), ("UA", "Ukraine"),
        ])
    else:
        try:
            with sqlite3.connect(DB_PATH, timeout=1) as _c:
                _c.row_factory = sqlite3.Row
                _r = _c.execute("SELECT country_code, country_name FROM ip_geo_cache WHERE ip = ?", (ip,)).fetchone()
                if _r:
                    cc, cn = _r["country_code"], _r["country_name"]
                else:
                    # Fire background thread for API lookup — never blocks the request
                    def _bg_geo(pip):
                        try:
                            import requests as _req
                            r2 = _req.get(
                                f"http://ip-api.com/json/{pip}?fields=status,country,countryCode,lat,lon",
                                timeout=3
                            )
                            if r2.status_code == 200:
                                d = r2.json()
                                if d.get("status") == "success":
                                    with sqlite3.connect(DB_PATH, timeout=3) as c2:
                                        c2.execute(
                                            "INSERT OR IGNORE INTO ip_geo_cache (ip, country_code, country_name, latitude, longitude) VALUES (?,?,?,?,?)",
                                            (pip, d.get("countryCode", "??"), d.get("country", "Unknown"), d.get("lat", 0), d.get("lon", 0))
                                        )
                                        c2.commit()
                        except Exception:
                            pass
                    threading.Thread(target=_bg_geo, args=(ip,), daemon=True).start()
        except Exception:
            pass

    entry = {
        "timestamp":    datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "ip":           ip,
        "method":       request.method,
        "path":         request.path,
        "query":        request.query_string.decode("utf-8", errors="ignore"),
        "attack_type":  attack_type,
        "severity":     severity,
        "action":       action,
        "payload":      payload_excerpt or "",
        "risk_score":   score,
        "risk_level":   level,
        "country_code": cc,
        "country_name": cn,
    }

    write_log(entry)
    g.waf_action = action
    g.waf_entry  = entry

    if action == "BLOCKED":
        return render_template("blocked.html", attack_type=attack_type, ip=ip), 403

# ─────────────────────────────────────────────
#  Target Routes
# ─────────────────────────────────────────────

@app.route("/")
def dashboard():
    return render_template("index.html")

@app.route("/index.html")
@app.route("/about.html")
@app.route("/contact.html")
def static_pages():
    return render_template("pages.html", page=request.path.lstrip("/"))

@app.route("/search")
def search():
    return jsonify({"status": "ok", "query": request.args.get("q", request.args.get("id", "")), "results": []})

@app.route("/comment", methods=["POST"])
def comment():
    return jsonify({"status": "ok", "comment": request.form.get("body", "")})

@app.route("/view")
def view():
    return jsonify({"status": "ok", "file": request.args.get("file", "")})

@app.route("/exec")
def execute():
    return jsonify({"status": "ok", "cmd": request.args.get("cmd", "")})

@app.route("/login")
def login():
    return jsonify({"status": "ok", "user": request.args.get("user", ""), "authenticated": False})

# ─────────────────────────────────────────────
#  API — Live Logs
# ─────────────────────────────────────────────

@app.route("/api/logs")
def api_logs():
    limit = int(request.args.get("limit", 100))
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        rows = conn.execute(
            "SELECT * FROM attack_logs ORDER BY id DESC LIMIT ?", (limit,)
        ).fetchall()
    return jsonify([dict(r) for r in rows])

# ─────────────────────────────────────────────
#  API — Stats
# ─────────────────────────────────────────────

@app.route("/api/stats")
def api_stats():
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row

        # Totals
        row = conn.execute("""
            SELECT
                COUNT(*)                                         AS total,
                SUM(CASE WHEN action='BLOCKED' THEN 1 ELSE 0 END) AS blocked,
                SUM(CASE WHEN action='FLAGGED' THEN 1 ELSE 0 END) AS flagged,
                SUM(CASE WHEN action='ALLOWED' THEN 1 ELSE 0 END) AS allowed
            FROM attack_logs
        """).fetchone()
        total, blocked, flagged, allowed = (
            row["total"] or 0, row["blocked"] or 0,
            row["flagged"] or 0, row["allowed"] or 0
        )

        # Attack distribution (exclude Normal)
        dist_rows = conn.execute("""
            SELECT attack_type, COUNT(*) AS cnt
            FROM attack_logs
            WHERE attack_type != 'Normal'
            GROUP BY attack_type
            ORDER BY cnt DESC
        """).fetchall()
        dist = {r["attack_type"]: r["cnt"] for r in dist_rows}

        # Top attacker IPs
        top_rows = conn.execute("""
            SELECT ip,
                   COUNT(*) AS cnt,
                   SUM(CASE WHEN action='BLOCKED' THEN 1 ELSE 0 END) AS blocked,
                   SUM(CASE WHEN action='FLAGGED' THEN 1 ELSE 0 END) AS flagged
            FROM attack_logs
            WHERE action IN ('BLOCKED','FLAGGED')
            GROUP BY ip
            ORDER BY cnt DESC
            LIMIT 5
        """).fetchall()
        top_ips = [{"ip": r["ip"], "count": r["cnt"],
                    "blocked": r["blocked"], "flagged": r["flagged"]} for r in top_rows]

    detection_rate = round((blocked + flagged) / max(total, 1) * 100, 1)

    return jsonify({
        "total":              total,
        "blocked":            blocked,
        "flagged":            flagged,
        "allowed":            allowed,
        "detection_rate":     detection_rate,
        "attack_distribution": dist,
        "top_attackers":      top_ips,
    })

# ─────────────────────────────────────────────
#  API — Daily Chart (attacks per day)
# ─────────────────────────────────────────────

@app.route("/api/daily")
def api_daily():
    days = int(request.args.get("days", 7))
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        rows = conn.execute("""
            SELECT
                substr(timestamp, 1, 10) AS day,
                SUM(CASE WHEN action='BLOCKED' THEN 1 ELSE 0 END) AS blocked,
                SUM(CASE WHEN action='FLAGGED' THEN 1 ELSE 0 END) AS flagged,
                SUM(CASE WHEN action='ALLOWED' THEN 1 ELSE 0 END) AS allowed,
                COUNT(*) AS total
            FROM attack_logs
            WHERE timestamp >= date('now', ?)
            GROUP BY day
            ORDER BY day ASC
        """, (f"-{days} days",)).fetchall()

    # Fill in missing days with zeros
    result = {}
    for i in range(days):
        d = (datetime.now() - timedelta(days=days-1-i)).strftime("%Y-%m-%d")
        result[d] = {"day": d, "blocked": 0, "flagged": 0, "allowed": 0, "total": 0}
    for r in rows:
        result[r["day"]] = dict(r)

    return jsonify(list(result.values()))

# ─────────────────────────────────────────────
#  API — IP Reputation
# ─────────────────────────────────────────────

# ─────────────────────────────────────────────
#  API — Risk Distribution
# ─────────────────────────────────────────────

@app.route("/api/risk")
def api_risk():
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row

        # Risk level distribution counts
        dist = conn.execute("""
            SELECT risk_level,
                   COUNT(*) AS count,
                   AVG(risk_score) AS avg_score,
                   MAX(risk_score) AS max_score
            FROM attack_logs
            GROUP BY risk_level
        """).fetchall()

        # Overall risk stats
        overall = conn.execute("""
            SELECT AVG(risk_score) AS avg_score,
                   MAX(risk_score) AS max_score,
                   COUNT(*) AS total
            FROM attack_logs
        """).fetchone()

        # Top risky IPs with avg score
        risky_ips = conn.execute("""
            SELECT ip,
                   ROUND(AVG(risk_score), 1) AS avg_score,
                   MAX(risk_score) AS max_score,
                   COUNT(*) AS total_requests,
                   SUM(CASE WHEN risk_level='critical' THEN 1 ELSE 0 END) AS critical_count,
                   SUM(CASE WHEN risk_level='high' THEN 1 ELSE 0 END) AS high_count
            FROM attack_logs
            GROUP BY ip
            ORDER BY avg_score DESC
            LIMIT 8
        """).fetchall()

        # Recent high-risk events
        recent_high = conn.execute("""
            SELECT id, timestamp, ip, attack_type, risk_score, risk_level, action
            FROM attack_logs
            WHERE risk_level IN ('critical','high')
            ORDER BY id DESC
            LIMIT 10
        """).fetchall()

    dist_map = {r["risk_level"]: {
        "count": r["count"],
        "avg_score": round(r["avg_score"] or 0, 1),
        "max_score": r["max_score"] or 0
    } for r in dist}

    # Ensure all levels present
    for lvl in ("critical", "high", "medium", "low"):
        dist_map.setdefault(lvl, {"count": 0, "avg_score": 0, "max_score": 0})

    return jsonify({
        "distribution":    dist_map,
        "overall_avg":     round(overall["avg_score"] or 0, 1),
        "overall_max":     overall["max_score"] or 0,
        "total":           overall["total"] or 0,
        "risky_ips":       [dict(r) for r in risky_ips],
        "recent_high_risk":[dict(r) for r in recent_high],
    })


@app.route("/api/geo")
def api_geo():
    """Return country-level attack distribution."""
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        
        # Aggregated attacks by country (excluding Normal traffic)
        rows = conn.execute("""
            SELECT country_code, country_name, COUNT(*) AS count,
                   MAX(risk_score) AS max_risk,
                   SUM(CASE WHEN action='BLOCKED' THEN 1 ELSE 0 END) AS blocked
            FROM attack_logs
            WHERE attack_type != 'Normal'
            GROUP BY country_code
            ORDER BY count DESC
        """).fetchall()
        
        # Recent geo-events feed
        recent = conn.execute("""
            SELECT country_code, country_name, attack_type, timestamp, ip
            FROM attack_logs
            WHERE attack_type != 'Normal'
            ORDER BY id DESC
            LIMIT 5
        """).fetchall()

    return jsonify({
        "distribution": [dict(r) for r in rows],
        "recent_geo":   [dict(r) for r in recent],
        "total_countries": len(rows)
    })

@app.route("/api/users-by-risk")
def api_users_by_risk():
    """Return unique IPs grouped by their worst (highest) risk tier."""
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row

        # Per-IP: worst risk level seen, avg score, total requests, top attack
        rows = conn.execute("""
            SELECT
                ip,
                ROUND(AVG(risk_score), 1)   AS avg_score,
                MAX(risk_score)              AS max_score,
                COUNT(*)                     AS total_requests,
                SUM(CASE WHEN action='BLOCKED' THEN 1 ELSE 0 END) AS blocked,
                SUM(CASE WHEN action='FLAGGED' THEN 1 ELSE 0 END) AS flagged,
                -- dominant risk level = level of their max risk score row
                (SELECT risk_level FROM attack_logs b
                 WHERE b.ip = a.ip ORDER BY risk_score DESC LIMIT 1) AS worst_level,
                (SELECT attack_type FROM attack_logs b
                 WHERE b.ip = a.ip AND attack_type != 'Normal'
                 GROUP BY attack_type ORDER BY COUNT(*) DESC LIMIT 1) AS top_attack
            FROM attack_logs a
            GROUP BY ip
        """).fetchall()

    # Group by worst risk level
    tiers = {"critical": [], "high": [], "medium": [], "low": []}
    for r in rows:
        lvl = r["worst_level"] or "low"
        tiers.setdefault(lvl, []).append({
            "ip":             r["ip"],
            "avg_score":      r["avg_score"] or 0,
            "max_score":      r["max_score"] or 0,
            "total_requests": r["total_requests"],
            "blocked":        r["blocked"],
            "flagged":        r["flagged"],
            "top_attack":     r["top_attack"] or "Normal",
        })

    # Sort each tier by avg_score desc, cap at 10
    for lvl in tiers:
        tiers[lvl].sort(key=lambda x: x["avg_score"], reverse=True)
        tiers[lvl] = tiers[lvl][:10]

    summary = {
        lvl: {"count": len(tiers[lvl]), "ips": tiers[lvl]}
        for lvl in ("critical", "high", "medium", "low")
    }
    return jsonify({"tiers": summary, "total_unique_ips": len(rows)})


# ─────────────────────────────────────────────
#  API — IP Reputation & Bans
# ─────────────────────────────────────────────

@app.route("/api/reputation")
def api_reputation():
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        rows = conn.execute("""
            SELECT * FROM ip_reputation
            ORDER BY blocked_count DESC
            LIMIT 20
        """).fetchall()
    return jsonify([dict(r) for r in rows])

@app.route("/api/bans")
def api_bans():
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        rows = conn.execute("SELECT * FROM banned_ips ORDER BY timestamp DESC").fetchall()
    return jsonify([dict(r) for r in rows])

@app.route("/api/bans/lift", methods=["POST"])
def api_bans_lift():
    ip = request.json.get("ip")
    if not ip:
        return jsonify({"error": "IP required"}), 400
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("DELETE FROM banned_ips WHERE ip = ?", (ip,))
        conn.commit()
    return jsonify({"status": "lifted", "ip": ip})

# ─────────────────────────────────────────────
#  API — Security Report
# ─────────────────────────────────────────────

def get_report_data():
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row

        summary = conn.execute("""
            SELECT
                COUNT(*)                                               AS total_requests,
                SUM(CASE WHEN action='BLOCKED' THEN 1 ELSE 0 END)     AS total_blocked,
                SUM(CASE WHEN action='FLAGGED' THEN 1 ELSE 0 END)     AS total_flagged,
                SUM(CASE WHEN action='ALLOWED' THEN 1 ELSE 0 END)     AS total_allowed,
                COUNT(DISTINCT ip)                                     AS unique_ips,
                MIN(timestamp)                                         AS first_event,
                MAX(timestamp)                                         AS last_event
            FROM attack_logs
        """).fetchone()

        attack_breakdown = conn.execute("""
            SELECT attack_type, severity,
                   COUNT(*) AS count,
                   COUNT(DISTINCT ip) AS unique_ips
            FROM attack_logs
            WHERE attack_type != 'Normal'
            GROUP BY attack_type
            ORDER BY count DESC
        """).fetchall()

        repeat_offenders = conn.execute("""
            SELECT ip, total_requests, blocked_count, flagged_count,
                   allowed_count, first_seen, last_seen
            FROM ip_reputation
            WHERE blocked_count + flagged_count > 0
            ORDER BY blocked_count DESC
            LIMIT 10
        """).fetchall()

        hourly = conn.execute("""
            SELECT substr(timestamp, 12, 2) AS hour,
                   COUNT(*) AS total,
                   SUM(CASE WHEN action='BLOCKED' THEN 1 ELSE 0 END) AS blocked
            FROM attack_logs
            GROUP BY hour
            ORDER BY hour
        """).fetchall()

    return {
        "summary":           dict(summary),
        "attack_breakdown":  [dict(r) for r in attack_breakdown],
        "repeat_offenders":  [dict(r) for r in repeat_offenders],
        "hourly_activity":   [dict(r) for r in hourly],
        "generated_at":      datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }

def get_recommendations(data):
    recs = []
    types = [t['attack_type'] for t in data['attack_breakdown']]
    if "SQL Injection" in types:
        recs.append("Enable parameterized queries and stronger input validation for database queries.")
    if "Local File Inclusion (LFI)" in types or "Directory Traversal" in types:
        recs.append("Ensure local file operations are strictly sanitized and sandboxed. Disable absolute path referencing.")
    if "Cross-Site Scripting (XSS)" in types:
        recs.append("Implement strict Content Security Policy (CSP) and sanitize all user inputs before rendering.")
    if "Command Injection" in types:
        recs.append("Never pass raw user input to system shells. Use whitelisted commands and safe APIs instead.")
    
    if data['summary']['total_blocked'] and data['summary']['total_blocked'] > data['summary']['total_allowed'] * 2:
        recs.append("High block-rate detected. Consider implementing a CAPTCHA on critical endpoints to deter automated scanners.")
    if not recs:
        recs.append("Maintain routine security audits and keep dependencies updated.")
    return recs

@app.route("/api/report")
def api_report():
    return jsonify(get_report_data())

# ─────────────────────────────────────────────
#  API — Auto Report Generator (PDF, DOCX, TXT, etc)
# ─────────────────────────────────────────────

@app.route("/api/report/download")
def download_report():
    fmt = request.args.get("format", "txt").lower()
    data = get_report_data()
    recs = get_recommendations(data)

    if fmt == "txt":
        lines = [
            "==========================================",
            "      SENTINELSHIELD SECURITY REPORT      ",
            "==========================================\n",
            f"Generated At: {data['generated_at']}\n",
            "--- 1. Attack Summary ---",
            f"Total Requests: {data['summary']['total_requests'] or 0}",
            f"Blocked Requests: {data['summary']['total_blocked'] or 0}",
            f"Safe/Allowed Requests: {data['summary']['total_allowed'] or 0}",
            f"Unique IP Addresses seen: {data['summary']['unique_ips'] or 0}\n",
            "--- 2. Attack Type Distribution ---"
        ]
        for a in data['attack_breakdown']:
            lines.append(f"{a['attack_type']} (Sev: {a['severity']}) -> {a['count']} attacks")
        
        lines.append("\n--- 3. Most Suspicious IPs ---")
        for i in data['repeat_offenders']:
            lines.append(f"{i['ip']} -> {i['blocked_count']} blocks (Total Reqs: {i['total_requests']})")
            
        lines.append("\n--- 4. Security Recommendations ---")
        for r in recs:
            lines.append(f"- {r}")

        return Response("\n".join(lines), mimetype="text/plain", headers={"Content-disposition": "attachment; filename=security_report.txt"})

    elif fmt == "csv":
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["Metric", "Value"])
        writer.writerow(["Total Requests", data['summary']['total_requests'] or 0])
        writer.writerow(["Blocked Requests", data['summary']['total_blocked'] or 0])
        writer.writerow(["Generated At", data['generated_at']])
        writer.writerow([])
        writer.writerow(["Attack Type", "Severity", "Count", "Unique IPs Targeting"])
        for a in data['attack_breakdown']:
            writer.writerow([a['attack_type'], a['severity'], a['count'], a['unique_ips']])
        return Response(output.getvalue(), mimetype="text/csv", headers={"Content-disposition": "attachment; filename=security_report.csv"})

    elif fmt == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", style="B", size=16)
        pdf.cell(200, 10, "SentinelShield Security Report", ln=True, align='C')
        pdf.set_font("Helvetica", size=10)
        pdf.cell(200, 10, f"Generated At: {data['generated_at']}", ln=True, align='C')
        
        pdf.ln(10)
        pdf.set_font("Helvetica", style="B", size=12)
        pdf.cell(200, 10, "1. Attack Summary", ln=True)
        pdf.set_font("Helvetica", size=10)
        pdf.cell(200, 8, f"Total Requests: {data['summary']['total_requests'] or 0}", ln=True)
        pdf.cell(200, 8, f"Blocked Requests: {data['summary']['total_blocked'] or 0}", ln=True)
        pdf.cell(200, 8, f"Safe/Allowed Requests: {data['summary']['total_allowed'] or 0}", ln=True)
        
        pdf.ln(5)
        pdf.set_font("Helvetica", style="B", size=12)
        pdf.cell(200, 10, "2. Attack Type Distribution", ln=True)
        pdf.set_font("Helvetica", size=10)
        for a in data['attack_breakdown']:
            pdf.cell(200, 8, f"{a['attack_type']} ({a['severity']}) -> {a['count']} attacks", ln=True)

        pdf.ln(5)
        pdf.set_font("Helvetica", style="B", size=12)
        pdf.cell(200, 10, "3. Most Suspicious IPs", ln=True)
        pdf.set_font("Helvetica", size=10)
        for i in data['repeat_offenders']:
            pdf.cell(200, 8, f"{i['ip']} -> {i['blocked_count']} blocks", ln=True)
            
        pdf.ln(10)
        pdf.set_font("Helvetica", style="B", size=12)
        pdf.cell(200, 10, "4. Security Recommendations", ln=True)
        pdf.set_font("Helvetica", size=10)
        for r in recs:
            pdf.cell(0, 8, f"- {r}", ln=True)

        return Response(pdf.output(), mimetype="application/pdf", headers={"Content-disposition": "attachment; filename=security_report.pdf"})

    elif fmt == "docx":
        doc = Document()
        doc.add_heading('SentinelShield Security Report', 0)
        doc.add_paragraph(f"Generated At: {data['generated_at']}")
        
        doc.add_heading('1. Attack Summary', level=1)
        doc.add_paragraph(f"Total Requests: {data['summary']['total_requests'] or 0}\nBlocked Requests: {data['summary']['total_blocked'] or 0}\nSafe/Allowed: {data['summary']['total_allowed'] or 0}")
        
        doc.add_heading('2. Attack Type Distribution', level=1)
        for a in data['attack_breakdown']:
            doc.add_paragraph(f"{a['attack_type']} ({a['severity']}) -> {a['count']} attacks", style='ListBullet')
            
        doc.add_heading('3. Most Suspicious IPs', level=1)
        for i in data['repeat_offenders']:
            doc.add_paragraph(f"{i['ip']} -> {i['blocked_count']} blocks", style='ListBullet')

        doc.add_heading('4. Security Recommendations', level=1)
        for r in recs:
            doc.add_paragraph(r, style='List Bullet')
            
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return Response(f.read(), mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-disposition": "attachment; filename=security_report.docx"})

    elif fmt == "html":
        return render_template("report_export.html", data=data, recs=recs)

    return jsonify({"error": "Unsupported format"}), 400

# ─────────────────────────────────────────────
#  API — Raw Attack Log Export (Legacy)
# ─────────────────────────────────────────────


@app.route("/api/export/csv")
def export_csv():
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        rows = conn.execute(
            "SELECT * FROM attack_logs ORDER BY id DESC"
        ).fetchall()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["ID","Timestamp","IP","Method","Path","Query","Attack Type","Severity","Action","Payload"])
    for r in rows:
        writer.writerow([r["id"], r["timestamp"], r["ip"], r["method"],
                         r["path"], r["query"], r["attack_type"],
                         r["severity"], r["action"], r["payload"]])

    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=sentinelshield_logs.csv"}
    )

# ─────────────────────────────────────────────
#  API — Clear
# ─────────────────────────────────────────────

@app.route("/api/clear", methods=["POST"])
def api_clear():
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("DELETE FROM attack_logs")
        conn.execute("DELETE FROM ip_reputation")
        conn.commit()
    with rate_lock:
        rate_tracker.clear()
    return jsonify({"status": "cleared"})

# ─────────────────────────────────────────────
#  Entrypoint
# ─────────────────────────────────────────────

if __name__ == "__main__":
    init_db()
    print("\n" + "═" * 60)
    print("  SentinelShield WAF/IDS  |  SQLite Backend Active")
    print(f"  DB: {DB_PATH}")
    print("  http://localhost:5000")
    print("═" * 60 + "\n")
    app.run(host="0.0.0.0", port=5000, debug=True)
